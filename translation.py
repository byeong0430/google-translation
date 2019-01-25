#!/usr/bin/python

import requests
import json
import docx # You also need to download python-docx module
from google.cloud import translate
import os, sys

# Variables
# Target language (List of supported languages: https://cloud.google.com/translate/docs/languages)
target_language = 'zh-TW'
input_dir = os.path.abspath('input/')
output_dir = os.path.abspath('output/')
input_file = [x for x in os.listdir(input_dir) if x.endswith('.docx')][0]

# Class
class Google:
    def __init__(self, translate):
        self.translate_client = translate.Client()

    def call_translation_api(self, text, target_language):
        tr = self.translate_client.translate(text, target_language)
        return tr['translatedText']

    def translate(self, *args):
        # Destructure array
        text, target_language = args

        # Remove white space of text
        text = text.strip()
        if not text:
            print('Input: "%s". \nYour input is an empty string. Next...' % text)
            return text

        result = self.call_translation_api(text, target_language)

        print('"%s" => "%s"' % (text, result))
        return result

# Initialise class
gg = Google(translate)

# Input docx
print('%s/%s' % (input_dir, input_file))
doc = docx.Document('%s/%s' % (input_dir, input_file))

# Header
for section in doc.sections:
    paragraphs = section.header.paragraphs
    for paragraph in paragraphs:
        if paragraph.text:
            translation = gg.translate(paragraph.text, target_language)
            paragraph.text = translation

doc.save('%s/translated_%s' % (output_dir, input_file))

# Paragraphs
for paragraph in doc.paragraphs:
    if paragraph.text:
        translation = gg.translate(paragraph.text, target_language)
        paragraph.text = translation

doc.save('%s/translated_%s' % (output_dir, input_file))

# Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text:
                translation = gg.translate(cell.text, target_language)
                cell.text = translation

doc.save('%s/translated_%s' % (output_dir, input_file))

# Footer
for section in doc.sections:
    paragraphs = section.footer.paragraphs
    for paragraph in paragraphs:
        if paragraph.text:
            translation = gg.translate(paragraph.text, target_language)
            paragraph.text = translation

doc.save('%s/translated_%s' % (output_dir, input_file))